import os
import sys
import gc
import time
import re
import shutil
import tempfile
import hashlib
import pickle
import warnings
from pathlib import Path
from typing import List, Dict, Tuple, Union, Optional, Any
from datetime import datetime
from dataclasses import dataclass, field
from typing import List
from dataclasses import dataclass
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Core libraries
import numpy as np
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# PDF processing
import fitz  # PyMuPDF
import easyocr

# NLP libraries
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords

# ML libraries
try:
    import torch
    from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, logging as transformers_logging
    transformers_logging.set_verbosity_error()
    torch_available = True
except ImportError:
    torch_available = False
    logger.warning("PyTorch/Transformers not available")

try:
    from sentence_transformers import SentenceTransformer, CrossEncoder
    import faiss
    sentence_transformers_available = True
except ImportError:
    sentence_transformers_available = False
    logger.warning("Sentence transformers/FAISS not available")

# Document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    docx_available = True
except ImportError:
    docx_available = False
    logger.warning("python-docx not available")

# Gradio
import gradio as gr

warnings.filterwarnings("ignore")

# =============================================================================
# CONFIGURATION
# =============================================================================

@dataclass
class EnhancedConfig:
    # Model configurations
    EMBEDDING_MODEL: str = "sentence-transformers/all-MiniLM-L6-v2"  # Smaller model for deployment
    RERANKER_MODEL: str = "cross-encoder/ms-marco-MiniLM-L-4-v2"  # Smaller reranker
    QA_MODEL: str = "google/flan-t5-base"  # Smaller T5 model
    
    # System limits
    MAX_PDFS: int = 10
    MAX_PDF_SIZE_MB: int = 50  # Reduced for deployment
    MAX_CHUNKS_RETRIEVE: int = 50
    TOP_K_AFTER_RERANK: int = 8
    
    # Processing parameters
    CHUNK_SIZE: int = 500
    CHUNK_OVERLAP: int = 50
    MIN_CHUNK_LENGTH: int = 50
    
    # OCR settings
    OCR_LANGUAGES: List[str] = field(default_factory=lambda: ["en"])
    USE_GPU: bool = False  # CPU only for deployment
    
    # Cache settings
    CACHE_DIR: str = "/tmp/qa_cache"
    USE_CACHE: bool = True

config = EnhancedConfig()

# =============================================================================
# NLTK SETUP
# =============================================================================

def setup_nltk():
    """Download required NLTK data"""
    try:
        nltk.download('punkt', quiet=True)
        nltk.download('stopwords', quiet=True)
        nltk.download('averaged_perceptron_tagger', quiet=True)
        logger.info("✅ NLTK setup complete")
        return True
    except Exception as e:
        logger.error(f"NLTK setup failed: {e}")
        return False

# =============================================================================
# OCR PROCESSOR
# =============================================================================

class EasyOCRProcessor:
    def __init__(self):
        self.reader = None
        self.initialized = False
        
    def initialize(self):
        """Initialize EasyOCR reader"""
        if self.initialized:
            return True
            
        try:
            logger.info("🔄 Initializing EasyOCR...")
            self.reader = easyocr.Reader(
                config.OCR_LANGUAGES,
                gpu=config.USE_GPU,
                model_storage_directory='/tmp/easyocr_models',
                download_enabled=True
            )
            self.initialized = True
            logger.info(f"✅ EasyOCR initialized (GPU: {config.USE_GPU})")
            return True
        except Exception as e:
            logger.error(f"Failed to initialize EasyOCR: {e}")
            return False
    
    def extract_text_from_image(self, image_path: str) -> str:
        """Extract text from image using EasyOCR"""
        if not self.initialized:
            if not self.initialize():
                return ""
        
        try:
            results = self.reader.readtext(image_path, detail=0)
            return " ".join(results) if results else ""
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return ""

# =============================================================================
# TEXT PROCESSOR
# =============================================================================

class EnhancedTextProcessor:
    def __init__(self):
        self.ocr_processor = EasyOCRProcessor()
        
    def extract_text_from_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """Extract text from PDF using PyMuPDF with OCR fallback"""
        try:
            doc = fitz.open(pdf_path)
            pages_data = []
            total_text = ""
            has_ocr_content = False
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Extract text
                text = page.get_text()
                
                # If no text found, try OCR
                ocr_text = ""
                if len(text.strip()) < 50:  # Likely image-based page
                    try:
                        # Convert page to image
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom
                        img_path = f"/tmp/page_{page_num}.png"
                        pix.save(img_path)
                        
                        # Apply OCR
                        ocr_text = self.ocr_processor.extract_text_from_image(img_path)
                        if ocr_text:
                            has_ocr_content = True
                            text = ocr_text
                        
                        # Cleanup
                        if os.path.exists(img_path):
                            os.remove(img_path)
                            
                    except Exception as e:
                        logger.warning(f"OCR failed for page {page_num}: {e}")
                
                pages_data.append({
                    "page_number": page_num + 1,
                    "text": text,
                    "has_ocr": bool(ocr_text),
                    "char_count": len(text)
                })
                
                total_text += f"\n[Page {page_num + 1}]\n{text}\n"
            
            doc.close()
            
            return {
                "text": total_text,
                "pages": pages_data,
                "total_pages": len(doc),
                "has_ocr_content": has_ocr_content,
                "total_chars": len(total_text)
            }
            
        except Exception as e:
            logger.error(f"Failed to process PDF {pdf_path}: {e}")
            raise e
    
    def create_chunks(self, text: str, doc_name: str) -> List[Dict]:
        """Create text chunks"""
        try:
            sentences = sent_tokenize(text)
            chunks = []
            current_chunk = []
            current_length = 0
            
            for sentence in sentences:
                sentence_length = len(sentence)
                
                if current_length + sentence_length > config.CHUNK_SIZE and current_chunk:
                    chunk_text = " ".join(current_chunk)
                    if len(chunk_text) >= config.MIN_CHUNK_LENGTH:
                        chunks.append({
                            "text": chunk_text,
                            "document": doc_name,
                            "chunk_id": len(chunks),
                            "length": len(chunk_text)
                        })
                    current_chunk = current_chunk[-1:] if config.CHUNK_OVERLAP > 0 else []
                    current_length = sum(len(s) for s in current_chunk)
                
                current_chunk.append(sentence)
                current_length += sentence_length
            
            # Add final chunk
            if current_chunk:
                chunk_text = " ".join(current_chunk)
                if len(chunk_text) >= config.MIN_CHUNK_LENGTH:
                    chunks.append({
                        "text": chunk_text,
                        "document": doc_name,
                        "chunk_id": len(chunks),
                        "length": len(chunk_text)
                    })
            
            return chunks
            
        except Exception as e:
            logger.error(f"Chunking failed: {e}")
            return []

# =============================================================================
# RETRIEVAL SYSTEM
# =============================================================================

class EnhancedRetriever:
    def __init__(self):
        self.embedding_model = None
        self.tfidf_vectorizer = None
        self.tfidf_matrix = None
        self.chunks = []
        self.embeddings = None
        self.faiss_index = None
        
    def initialize_models(self):
        """Initialize retrieval models"""
        try:
            if sentence_transformers_available:
                logger.info(f"Loading embedding model: {config.EMBEDDING_MODEL}")
                self.embedding_model = SentenceTransformer(config.EMBEDDING_MODEL)
                logger.info("✅ Embedding model loaded")
            else:
                logger.warning("Sentence transformers not available")
                
        except Exception as e:
            logger.error(f"Failed to initialize retrieval models: {e}")
    
    def index_documents(self, chunks: List[Dict]):
        """Create search index"""
        try:
            self.chunks = chunks
            texts = [chunk["text"] for chunk in chunks]
            
            # Create TF-IDF index
            self.tfidf_vectorizer = TfidfVectorizer(
                max_features=5000,
                stop_words="english",
                ngram_range=(1, 2),
                min_df=1,
                max_df=0.95
            )
            self.tfidf_matrix = self.tfidf_vectorizer.fit_transform(texts)
            
            # Create semantic embeddings if available
            if self.embedding_model:
                self.embeddings = self.embedding_model.encode(
                    texts,
                    convert_to_numpy=True,
                    show_progress_bar=False,
                    batch_size=8
                )
                
                # Create FAISS index
                if sentence_transformers_available and self.embeddings is not None:
                    dimension = self.embeddings.shape[1]
                    self.faiss_index = faiss.IndexFlatIP(dimension)
                    self.faiss_index.add(self.embeddings.astype('float32'))
            
            logger.info(f"✅ Indexed {len(chunks)} chunks")
            
        except Exception as e:
            logger.error(f"Indexing failed: {e}")
    
    def retrieve_chunks(self, query: str, top_k: int = config.MAX_CHUNKS_RETRIEVE) -> List[Dict]:
        """Retrieve relevant chunks"""
        try:
            if not self.chunks:
                return []
            
            # TF-IDF search
            lexical_scores = np.zeros(len(self.chunks))
            if self.tfidf_vectorizer and self.tfidf_matrix is not None:
                query_tfidf = self.tfidf_vectorizer.transform([query])
                lexical_scores = cosine_similarity(query_tfidf, self.tfidf_matrix)[0]
            
            # Semantic search
            semantic_scores = np.zeros(len(self.chunks))
            if self.embedding_model and self.embeddings is not None:
                query_embedding = self.embedding_model.encode([query], convert_to_numpy=True)
                if self.faiss_index:
                    similarities, indices = self.faiss_index.search(
                        query_embedding.astype('float32'), 
                        min(len(self.chunks), top_k)
                    )
                    for idx, sim in zip(indices[0], similarities[0]):
                        if idx < len(semantic_scores):
                            semantic_scores[idx] = sim
                else:
                    semantic_scores = cosine_similarity(query_embedding, self.embeddings)[0]
            
            # Combine scores
            combined_scores = 0.3 * lexical_scores + 0.7 * semantic_scores
            
            # Get top results
            top_indices = np.argsort(combined_scores)[::-1][:top_k]
            
            results = []
            for idx in top_indices:
                if idx < len(self.chunks) and combined_scores[idx] > 0:
                    chunk = dict(self.chunks[idx])
                    chunk["score"] = float(combined_scores[idx])
                    results.append(chunk)
            
            return results
            
        except Exception as e:
            logger.error(f"Retrieval failed: {e}")
            return []

# =============================================================================
# ANSWER GENERATOR
# =============================================================================

class EnhancedAnswerGenerator:
    def __init__(self):
        self.tokenizer = None
        self.model = None
        self.reranker = None
        
    def initialize_models(self):
        """Initialize generation models"""
        try:
            if torch_available:
                logger.info(f"Loading QA model: {config.QA_MODEL}")
                self.tokenizer = AutoTokenizer.from_pretrained(config.QA_MODEL)
                self.model = AutoModelForSeq2SeqLM.from_pretrained(config.QA_MODEL)
                logger.info("✅ QA model loaded")
            
            if sentence_transformers_available:
                logger.info(f"Loading reranker: {config.RERANKER_MODEL}")
                self.reranker = CrossEncoder(config.RERANKER_MODEL)
                logger.info("✅ Reranker loaded")
                
        except Exception as e:
            logger.error(f"Failed to initialize generation models: {e}")
    
    def rerank_chunks(self, query: str, chunks: List[Dict]) -> List[Dict]:
        """Rerank chunks using cross-encoder"""
        try:
            if not self.reranker or not chunks:
                return chunks[:config.TOP_K_AFTER_RERANK]
            
            pairs = [[query, chunk["text"]] for chunk in chunks]
            scores = self.reranker.predict(pairs)
            
            for chunk, score in zip(chunks, scores):
                chunk["rerank_score"] = float(score)
            
            reranked = sorted(chunks, key=lambda x: x["rerank_score"], reverse=True)
            return reranked[:config.TOP_K_AFTER_RERANK]
            
        except Exception as e:
            logger.error(f"Reranking failed: {e}")
            return chunks[:config.TOP_K_AFTER_RERANK]
    
    def generate_answer(self, question: str, chunks: List[Dict]) -> Dict[str, Any]:
        """Generate answer from chunks"""
        try:
            if not chunks:
                return {
                    "answer": "I don't have enough information to answer this question.",
                    "confidence": 0.0,
                    "sources": []
                }
            
            # Prepare context
            context = " ".join([chunk["text"][:500] for chunk in chunks[:5]])
            
            # Generate answer
            if self.model and self.tokenizer:
                prompt = f"Question: {question}\nContext: {context}\nAnswer:"
                
                inputs = self.tokenizer(
                    prompt,
                    max_length=512,
                    truncation=True,
                    return_tensors="pt"
                )
                
                with torch.no_grad():
                    outputs = self.model.generate(
                        inputs["input_ids"],
                        max_length=200,
                        num_beams=3,
                        temperature=0.7,
                        do_sample=True,
                        pad_token_id=self.tokenizer.eos_token_id
                    )
                
                answer = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
                # Remove the prompt from answer
                if "Answer:" in answer:
                    answer = answer.split("Answer:")[-1].strip()
            else:
                # Extractive fallback
                sentences = []
                for chunk in chunks[:3]:
                    chunk_sentences = sent_tokenize(chunk["text"])
                    sentences.extend(chunk_sentences[:2])
                answer = " ".join(sentences[:5])
            
            # Create sources
            sources = []
            for i, chunk in enumerate(chunks[:5]):
                sources.append({
                    "id": i + 1,
                    "document_name": chunk.get("document", "Unknown"),
                    "text_preview": chunk["text"][:200] + "...",
                    "score": chunk.get("rerank_score", chunk.get("score", 0.0))
                })
            
            return {
                "answer": answer,
                "confidence": min(0.9, max(0.3, np.mean([s["score"] for s in sources]))),
                "sources": sources,
                "metadata": {
                    "chunks_used": len(chunks),
                    "context_length": len(context)
                }
            }
            
        except Exception as e:
            logger.error(f"Answer generation failed: {e}")
            return {
                "answer": f"Error generating answer: {str(e)}",
                "confidence": 0.0,
                "sources": []
            }

# =============================================================================
# WORD DOCUMENT GENERATOR
# =============================================================================

class WordDocumentGenerator:
    def __init__(self):
        self.temp_dir = "/tmp"
        
    def generate_filename(self, question: str, pdf_names: List[str]) -> str:
        """Generate custom filename: Question (PDF name).docx"""
        # Clean question
        clean_question = re.sub(r'[<>:"/\\|?*]', '', question.strip())
        clean_question = clean_question[:50] + "..." if len(clean_question) > 50 else clean_question
        
        # Process PDF names
        pdf_parts = []
        for name in pdf_names:
            clean_name = Path(name).stem if hasattr(Path(name), 'stem') else name.replace('.pdf', '')
            pdf_parts.append(clean_name)
        
        if len(pdf_parts) == 1:
            pdf_part = pdf_parts[0]
        else:
            pdf_part = " & ".join(pdf_parts[:2])
            if len(pdf_parts) > 2:
                pdf_part += f" +{len(pdf_parts)-2} more"
        
        filename = f"{clean_question} ({pdf_part}).docx"
        
        # Ensure reasonable length
        if len(filename) > 200:
            clean_question = clean_question[:80] + "..."
            filename = f"{clean_question} ({pdf_part}).docx"
        
        return filename
    
    def create_word_document(self, question: str, answer: str, sources: List[Dict], 
                           pdf_names: List[str]) -> Optional[str]:
        """Create Word document with custom filename"""
        try:
            if not docx_available:
                logger.error("python-docx not available")
                return None
            
            # Generate custom filename
            filename = self.generate_filename(question, pdf_names)
            doc_path = os.path.join(self.temp_dir, filename)
            
            # Create document
            doc = Document()
            
            # Title
            title = doc.add_heading('PDF Document Analysis', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Question
            doc.add_heading('Question:', level=1)
            doc.add_paragraph(question)
            
            # Answer
            doc.add_heading('Answer:', level=1)
            doc.add_paragraph(answer)
            
            # Sources
            if sources:
                doc.add_heading('Sources:', level=1)
                for i, source in enumerate(sources, 1):
                    doc.add_paragraph(
                        f"{i}. {source.get('document_name', 'Unknown Document')}\n"
                        f"   Preview: {source.get('text_preview', 'No preview available')}\n"
                        f"   Relevance Score: {source.get('score', 0.0):.3f}",
                        style='List Number'
                    )
            
            # Footer
            doc.add_paragraph()
            footer = doc.add_paragraph('Generated by Enhanced PDF QA System')
            footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Save document
            doc.save(doc_path)
            logger.info(f"✅ Word document created: {filename}")
            
            return doc_path
            
        except Exception as e:
            logger.error(f"Failed to create Word document: {e}")
            return None

# =============================================================================
# MAIN QA SYSTEM
# =============================================================================

class EnhancedPDFQASystem:
    def __init__(self):
        logger.info("🚀 Initializing Enhanced PDF QA System...")
        
        self.text_processor = EnhancedTextProcessor()
        self.retriever = EnhancedRetriever()
        self.answer_generator = EnhancedAnswerGenerator()
        self.word_generator = WordDocumentGenerator()
        
        self.documents = []
        self.all_chunks = []
        self.processed_files = []
        
        # Initialize models
        self.initialize_system()
        
    def initialize_system(self):
        """Initialize all system components"""
        try:
            # Setup NLTK
            setup_nltk()
            
            # Initialize models
            self.retriever.initialize_models()
            self.answer_generator.initialize_models()
            
            # Create cache directory
            os.makedirs(config.CACHE_DIR, exist_ok=True)
            
            logger.info("✅ System initialization complete!")
            
        except Exception as e:
            logger.error(f"System initialization failed: {e}")
    
    def process_pdfs(self, file_paths: List[str]) -> Dict[str, Any]:
        """Process uploaded PDF files"""
        try:
            logger.info(f"📚 Processing {len(file_paths)} PDF(s)...")
            
            all_chunks = []
            processed_docs = []
            
            for i, file_path in enumerate(file_paths):
                try:
                    logger.info(f"Processing PDF {i+1}/{len(file_paths)}: {Path(file_path).name}")
                    
                    # Extract text
                    pdf_data = self.text_processor.extract_text_from_pdf(file_path)
                    
                    # Create chunks
                    doc_name = Path(file_path).name
                    chunks = self.text_processor.create_chunks(pdf_data["text"], doc_name)
                    
                    all_chunks.extend(chunks)
                    processed_docs.append({
                        "name": doc_name,
                        "path": file_path,
                        "pages": pdf_data["total_pages"],
                        "has_ocr": pdf_data["has_ocr_content"],
                        "chunks": len(chunks)
                    })
                    
                    logger.info(f"✅ Processed {doc_name}: {pdf_data['total_pages']} pages, {len(chunks)} chunks")
                    
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
                    continue
            
            if not all_chunks:
                return {"success": False, "error": "No chunks were created from the PDFs"}
            
            # Index documents
            self.retriever.index_documents(all_chunks)
            
            # Store state
            self.all_chunks = all_chunks
            self.documents = processed_docs
            self.processed_files = file_paths
            
            # Create summary
            total_pages = sum(doc["pages"] for doc in processed_docs)
            ocr_docs = sum(1 for doc in processed_docs if doc["has_ocr"])
            
            summary = f"""**✅ Processing Complete!**

• **Documents processed:** {len(processed_docs)}
• **Total pages:** {total_pages}
• **Text chunks created:** {len(all_chunks)}
• **Documents with OCR content:** {ocr_docs}

**📋 Document Details:**
"""
            
            for i, doc in enumerate(processed_docs, 1):
                ocr_indicator = " 🔍(includes OCR)" if doc["has_ocr"] else ""
                summary += f"\n{i}. **{doc['name']}** - {doc['pages']} pages{ocr_indicator}"
            
            return {
                "success": True,
                "summary": summary,
                "suggestions": [
                    "Provide a comprehensive summary of the key concepts",
                    "What are the main methodologies discussed?",
                    "Analyze the key findings with examples",
                    "Compare different approaches mentioned in the documents"
                ]
            }
            
        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            return {"success": False, "error": str(e)}
    
    def answer_question(self, question: str, document_filter: Optional[List[str]] = None) -> Dict[str, Any]:
        """Answer a question based on processed documents"""
        try:
            if not self.all_chunks:
                return {
                    "answer": "Please upload and process PDF documents first.",
                    "confidence": 0.0,
                    "sources": [],
                    "suggestions": []
                }
            
            start_time = time.time()
            
            # Retrieve relevant chunks
            chunks = self.retriever.retrieve_chunks(question)
            
            if not chunks:
                return {
                    "answer": "No relevant information found for this question.",
                    "confidence": 0.0,
                    "sources": [],
                    "suggestions": []
                }
            
            # Rerank chunks
            reranked_chunks = self.answer_generator.rerank_chunks(question, chunks)
            
            # Generate answer
            result = self.answer_generator.generate_answer(question, reranked_chunks)
            
            # Create Word document
            word_doc_path = None
            if self.processed_files:
                pdf_names = [Path(f).name for f in self.processed_files]
                word_doc_path = self.word_generator.create_word_document(
                    question, result["answer"], result["sources"], pdf_names
                )
            
            processing_time = time.time() - start_time
            
            return {
                "answer": result["answer"],
                "confidence": result["confidence"],
                "sources": result["sources"],
                "processing_time": processing_time,
                "query_type": "general",
                "metadata": result.get("metadata", {}),
                "download_links": {"word": word_doc_path} if word_doc_path else {},
                "suggestions": [
                    "Can you elaborate on this topic with more examples?",
                    "What are the practical implications of this?",
                    "How does this compare to other approaches?",
                    "What are the key takeaways from this analysis?"
                ]
            }
            
        except Exception as e:
            logger.error(f"Question answering failed: {e}")
            return {
                "answer": f"Error processing question: {str(e)}",
                "confidence": 0.0,
                "sources": [],
                "suggestions": []
            }

# =============================================================================
# GRADIO INTERFACE
# =============================================================================

# Initialize the QA system
logger.info("Initializing QA System...")
qa_system = EnhancedPDFQASystem()

# Global variables
current_documents = []
processing_status = "Ready to process documents"
conversation_history = []

def validate_files(files) -> Tuple[bool, str]:
    """Validate uploaded files"""
    if not files:
        return False, "Please upload at least one PDF file"
    
    if len(files) > config.MAX_PDFS:
        return False, f"Maximum {config.MAX_PDFS} PDF files allowed"
    
    for file in files:
        if not file.name.lower().endswith('.pdf'):
            return False, f"Only PDF files allowed. '{file.name}' is not a PDF."
        
        try:
            file_size_mb = os.path.getsize(file.name) / (1024 * 1024)
            if file_size_mb > config.MAX_PDF_SIZE_MB:
                return False, f"File '{file.name}' is too large ({file_size_mb:.1f}MB). Max: {config.MAX_PDF_SIZE_MB}MB"
        except:
            pass
    
    return True, "Files validated successfully"

def process_documents(files, progress=gr.Progress()) -> Tuple[str, str, str]:
    """Process uploaded PDF documents"""
    global current_documents, processing_status
    
    try:
        progress(0.1, desc="Validating files...")
        is_valid, validation_msg = validate_files(files)
        
        if not is_valid:
            processing_status = f"❌ Validation Error: {validation_msg}"
            return processing_status, "", ""
        
        progress(0.2, desc="Preparing files...")
        file_paths = [file.name for file in files]
        
        progress(0.4, desc="Processing PDFs with OCR...")
        processing_status = "🔄 Processing documents (this may take a few minutes)..."
        
        result = qa_system.process_pdfs(file_paths)
        
        progress(0.9, desc="Finalizing...")
        
        if result["success"]:
            current_documents = file_paths
            processing_status = "✅ Documents processed successfully!"
            
            suggestions = "\n".join([f"• {s}" for s in result.get("suggestions", [])])
            suggestions_text = f"**💡 Suggested Questions:**\n\n{suggestions}"
            
            progress(1.0, desc="Complete!")
            return processing_status, result.get("summary", ""), suggestions_text
        else:
            processing_status = f"❌ Processing Error: {result.get('error', 'Unknown error')}"
            return processing_status, "", ""
            
    except Exception as e:
        processing_status = f"❌ Unexpected Error: {str(e)}"
        return processing_status, "", ""

def answer_question(question: str, document_filter: str, progress=gr.Progress()) -> Tuple[str, str, str, str]:
    """Answer a question based on processed documents"""
    global conversation_history
    
    if not question.strip():
        return "Please enter a question.", "", "", ""
    
    if not current_documents:
        return "Please upload and process PDF documents first.", "", "", ""
    
    try:
        progress(0.2, desc="Analyzing question...")
        progress(0.5, desc="Searching documents...")
        progress(0.8, desc="Generating answer...")
        
        result = qa_system.answer_question(question)
        
        # Format answer with metadata
        enhanced_answer = f"{result['answer']}\n\n"
        enhanced_answer += f"**📊 Response Metadata:**\n"
        enhanced_answer += f"• Confidence: {result['confidence']:.3f}\n"
        enhanced_answer += f"• Processing Time: {result['processing_time']:.2f}s\n"
        enhanced_answer += f"• Chunks Used: {result['metadata'].get('chunks_used', 0)}"
        
        # Format sources
        sources_text = "**📍 Sources:**\n\n"
        for source in result.get('sources', []):
            sources_text += f"{source['id']}. **{source['document_name']}**\n"
            sources_text += f"   Preview: {source['text_preview']}\n"
            sources_text += f"   Score: {source['score']:.3f}\n\n"
        
        # Format suggestions
        suggestions = "\n".join([f"• {s}" for s in result.get("suggestions", [])])
        suggestions_text = f"**💡 Follow-up Questions:**\n\n{suggestions}"
        
        # Format downloads
        downloads_text = "**📥 Download:**\n\n"
        if result.get("download_links", {}).get("word"):
            filename = os.path.basename(result["download_links"]["word"])
            downloads_text += f"📄 **Word Document:** `{filename}`\n"
            downloads_text += "*File saved in temporary directory*"
        else:
            downloads_text += "No download available"
        
        # Add to conversation history
        conversation_history.append({
            "question": question,
            "answer": result['answer'],
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        })
        
        progress(1.0, desc="Complete!")
        
        return enhanced_answer, sources_text, suggestions_text, downloads_text
        
    except Exception as e:
        return f"Error: {str(e)}", "", "", ""

def create_interface():
    """Create Gradio interface"""
    with gr.Blocks(
        theme=gr.themes.Soft(),
        title="Enhanced PDF QA System",
        css="""
        .gradio-container {
            font-family: 'Segoe UI', sans-serif !important;
        }
        .main-header {
            text-align: center;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 20px;
            border-radius: 10px;
            margin-bottom: 20px;
        }
        """
    ) as demo:
        
        # Header
        gr.HTML("""
        <div class="main-header">
            <h1>🔍 Enhanced PDF Document QA System</h1>
            <p>AI-powered question answering with OCR support and downloadable responses</p>
        </div>
        """)
        
        with gr.Row():
            with gr.Column(scale=2):
                # Upload section
                gr.HTML("<h2>📚 Upload PDF Documents</h2>")
                
                file_upload = gr.File(
                    label=f"Select PDF files (Max: {config.MAX_PDFS} files, {config.MAX_PDF_SIZE_MB}MB each)",
                    file_count="multiple",
                    file_types=[".pdf"]
                )
                
                process_btn = gr.Button("🚀 Process Documents", variant="primary")
                processing_status_display = gr.HTML(value="Ready to process documents")
                
                # Question section
                gr.HTML("<h2>❓ Ask Questions</h2>")
                
                question_input = gr.Textbox(
                    label="Your Question",
                    placeholder="Ask detailed questions about your documents...",
                    lines=3
                )
                
                document_filter = gr.Textbox(
                    label="Filter by Documents (optional)",
                    placeholder="Enter document names separated by commas",
                    lines=1
                )
                
                ask_btn = gr.Button("🤔 Ask Question", variant="secondary")
            
            with gr.Column(scale=1):
                gr.HTML("""
                <div style="background: #f8f9fa; padding: 15px; border-radius: 8px;">
                    <h3>🎯 System Features</h3>
                    <ul>
                        <li><strong>🔍 OCR Support</strong> - Reads text from images</li>
                        <li><strong>📚 Multi-Document</strong> - Query multiple PDFs</li>
                        <li><strong>📍 Source Attribution</strong> - Shows exact sources</li>
                        <li><strong>📥 Word Downloads</strong> - Custom filenames</li>
                        <li><strong>🤖 AI-Powered</strong> - Advanced NLP models</li>
                    </ul>
                </div>
                """)
                
                document_summary = gr.Markdown("Upload documents to see summary here.")
        
        # Results section
        gr.HTML("<h2>💡 Results</h2>")
        
        with gr.Tabs():
            with gr.Tab("📝 Answer"):
                answer_display = gr.Markdown("Answers will appear here.")
            
            with gr.Tab("📍 Sources"):
                sources_display = gr.Markdown("Source information will appear here.")
            
            with gr.Tab("💡 Suggestions"):
                suggestions_display = gr.Markdown("Question suggestions will appear here.")
            
            with gr.Tab("📥 Downloads"):
                downloads_display = gr.Markdown("Download links will appear here.")
        
        # Event handlers
        process_btn.click(
            fn=process_documents,
            inputs=[file_upload],
            outputs=[processing_status_display, document_summary, suggestions_display],
            show_progress=True
        )
        
        ask_btn.click(
            fn=answer_question,
            inputs=[question_input, document_filter],
            outputs=[answer_display, sources_display, suggestions_display, downloads_display],
            show_progress=True
        )
        
        question_input.submit(
            fn=answer_question,
            inputs=[question_input, document_filter],
            outputs=[answer_display, sources_display, suggestions_display, downloads_display]
        )
    
    return demo

# =============================================================================
# MAIN APPLICATION
# =============================================================================

if __name__ == "__main__":
    try:
        logger.info("🚀 Starting Enhanced PDF QA System...")
        
        # Create and launch interface
        demo = create_interface()
        
        # Launch with configuration for deployment
        demo.launch(
            server_name="0.0.0.0",  # Accept connections from any IP
            server_port=int(os.environ.get("PORT", 7860)),  # Use PORT env var or default
            share=False,
            show_error=True,
            inbrowser=False  # Don't auto-open browser in deployment
        )
        
    except Exception as e:
        logger.error(f"Failed to start application: {e}")
        raise e